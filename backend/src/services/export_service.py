import sys
import json
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def export_to_docx(questions, output_path):
    doc = Document()
    doc.add_heading('Generated Question Bank', 0)

    part_a = [q for q in questions if int(q.get('marks', 0)) == 2]
    part_b = [q for q in questions if int(q.get('marks', 0)) == 13]
    part_c = [q for q in questions if int(q.get('marks', 0)) == 16]

    sections = [
        ("Part-A (2 Marks)", part_a),
        ("Part-B (13 Marks)", part_b),
        ("Part-C (16 Marks)", part_c),
    ]

    for title, part_questions in sections:
        if not part_questions:
            continue
            
        doc.add_heading(title, level=1)
        
        for idx, q in enumerate(part_questions, 1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. ").bold = True
            p.add_run(q.get('text', ''))
            
            if q.get('modelAnswer'):
                ans_p = doc.add_paragraph()
                ans_p.add_run("Answer Key:").bold = True
                doc.add_paragraph(q.get('modelAnswer', ''))
            
            doc.add_paragraph() # spacing

    doc.save(output_path)

def export_to_pdf(questions, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    Story = []

    title_style = styles['Heading1']
    heading2_style = styles['Heading2']
    body_style = styles['BodyText']

    Story.append(Paragraph("Generated Question Bank", title_style))
    Story.append(Spacer(1, 12))

    part_a = [q for q in questions if int(q.get('marks', 0)) == 2]
    part_b = [q for q in questions if int(q.get('marks', 0)) == 13]
    part_c = [q for q in questions if int(q.get('marks', 0)) == 16]

    sections = [
        ("Part-A (2 Marks)", part_a),
        ("Part-B (13 Marks)", part_b),
        ("Part-C (16 Marks)", part_c),
    ]

    for title, part_questions in sections:
        if not part_questions:
            continue
            
        Story.append(Paragraph(title, heading2_style))
        Story.append(Spacer(1, 6))

        for idx, q in enumerate(part_questions, 1):
            q_text = q.get('text', '').replace('\n', '<br/>')
            text_html = f"<b>{idx}.</b> {q_text}"
            Story.append(Paragraph(text_html, body_style))
            
            if q.get('modelAnswer'):
                ans_text = q.get('modelAnswer', '').replace('\n', '<br/>')
                Story.append(Paragraph(f"<b>Answer Key:</b><br/>{ans_text}", body_style))

            Story.append(Spacer(1, 12))

    doc.build(Story)

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python export_service.py <format> <input_json> <output_path>")
        sys.exit(1)

    export_format = sys.argv[1].lower()
    input_json = sys.argv[2]
    output_path = sys.argv[3]

    if not os.path.exists(input_json):
        print(f"Error: Input file {input_json} not found.")
        sys.exit(1)

    with open(input_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    questions = data.get('questions', [])

    if export_format == 'docx':
        export_to_docx(questions, output_path)
    elif export_format == 'pdf':
        export_to_pdf(questions, output_path)
    else:
        print(f"Unsupported format: {export_format}")
        sys.exit(1)

    print(f"Successfully generated {output_path}")
